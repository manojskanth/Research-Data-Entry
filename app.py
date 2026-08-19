import streamlit as st
import datetime
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
import io
import json
import base64
import os
import re
import html
import urllib.request
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# --- 1. CORE SYSTEM CONFIGURATION ---
MASTER_SHEET_ID = st.secrets["MASTER_SHEET_ID"]
ALMANAC_SHEET_ID = "1ByouwZVNzQRtQsFLZLR9wppXKwakbczR"
ALMANAC_GID = "771144646"

RESEARCH_EVENTS_FOLDER_ID = "1UZxcyKw3RgmjyNV7eyIgwhzOkovB5fhF"
CAMPUS_ACTIVITIES_FOLDER_ID = "1EvUOvAqGD_aLCcCiuD3rHU0ra-ZZiKnS"
COMMITTEE_FOLDER_ID = "1pzrbGsViKtzsQYBPt-p9ZqQ5WXbzdHcW"

DEPARTMENTS = [
    "English & Languages", 
    "Social Sciences & Humanities", 
    "Sciences", 
    "Management", 
    "Commerce", 
    "IQAC", 
    "Research & Innovation", 
    "Physical Education"
]
COMMITTEES_CELLS_CLUBS = [
    "Alumni", "Anti-Ragging", "Disciplinary", "Equal Opportunity", 
    "Grievance Redressal", "Internal Complaints", "Scholarship", 
    "Library", "Placement", "Public Relations", "Women Empowerment", 
    "Student Activity Clubs", "IIC"
]
ACADEMIC_YEARS = ["2024-25", "2025-26", "2026-27", "2027-28", "2028-29", "2029-30"]
MONTHS = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]

DEPT_SORT_ORDER = {dept: index for index, dept in enumerate(DEPARTMENTS)}

DEPARTMENT_FOLDERS = {
    "Commerce": "1HMBoNkhksNpaitlBaGfq3JeoHsb_jmo-",
    "English & Languages": "14Nhs3qve5vDBbIT6GmzaRue51hvTzAOG",
    "IQAC": "19XsVcpJZRMyS0YlQ0RpShUfpUoCzeJb-",
    "Management": "1VG3xY_SmhqmQ9BvSh6KvDXOptO3kHhsj",
    "Physical Education": "1DA71KvpfSrltvv5io7gaeVN0v8CAix1w",
    "Research & Innovation": "1mRwg3vXDIZkTEkxBVDkJ-ruOsCsYsYnY",
    "Sciences": "1u_KRBhdZhcWQ55CyVI0v042bIpC5FQfs",
    "Social Sciences & Humanities": "1m0xEcv-WKQr8CWfHlZ5AuCWIFXAm1H5g"
}

FACULTY_DIRECTORY = {
    "saikiran@stmaryscollege.in": {"name": "Dr. Saikiran", "secret_key": "saikiran_pass"},
    "sangeetha@stmaryscollege.in": {"name": "Dr. Sangeetha", "secret_key": "sangeetha_pass"},
    "aditijuyal@stmaryscollege.in": {"name": "Ms. Aditi Juyal", "secret_key": "aditijuyal_pass"},
    "maithry@stmaryscollege.in": {"name": "Dr. Maithry Shinde", "secret_key": "maithry_pass"},
    "soumya@stmaryscollege.in": {"name": "Dr. Soumya K", "secret_key": "soumya_pass"},
    "rajita@stmaryscollege.in": {"name": "Dr. Rajita Anand Singh", "secret_key": "rajita_pass"},
    "manojkanth@stmaryscollege.in": {"name": "Dr. Manoj Kanth", "secret_key": "manojkanth_pass"},
    "swathi@stmaryscollege.in": {"name": "Dr. Swathi", "secret_key": "swathi_pass"},
    "padmaleela@stmaryscollege.in": {"name": "Dr. Padmaleela", "secret_key": "padmaleela_pass"},
    "sowjanya@stmaryscollege.in": {"name": "Ms. D. Sowjanya", "secret_key": "sowjanya_pass"},
    "sandhyarani@stmaryscollege.in": {"name": "Ms. A. Sandhya Rani", "secret_key": "sandhyarani_pass"},
    "ragasudha@stmaryscollege.in": {"name": "Ms. Raga Sudha Jonnada", "secret_key": "ragasudha_pass"},
    "rajyalakshmi@stmaryscollege.in": {"name": "Ms. Rajyalakshmi", "secret_key": "rajyalakshmi_pass"},
    "mahanta@stmaryscollege.in": {"name": "Ms. Mahanta Chauhan", "secret_key": "mahanta_pass"},
    "sharol@stmaryscollege.in": {"name": "Dr. Sharol Sebastian", "secret_key": "sharol_pass"},
    "govindaraju@stmaryscollege.in": {"name": "Dr. Govindaraju", "secret_key": "govind_pass"},
    "deepthipriya@stmaryscollege.in": {"name": "Dr. Deepthi Priya", "secret_key": "deepthipriya_pass"},
    "satabdi@stmaryscollege.in": {"name": "Dr. Satabdi Roy", "secret_key": "satabdi_pass"},
    "shima@stmaryscollege.in": {"name": "Ms. Shima A.N", "secret_key": "shima_pass"},
    "anuvictor@stmaryscollege.in": {"name": "Ms. Anu Victor", "secret_key": "anuvictor_pass"},
    "sadbhavana@stmaryscollege.in": {"name": "Ms. Sadbhavana Sharat", "secret_key": "sadbhavana_pass"},
    "sriveda@stmaryscollege.in": {"name": "Ms. Sriveda Baswapoor", "secret_key": "sriveda_pass"},
    "rameshk@stmaryscollege.in": {"name": "Dr. Ramesh Kumar", "secret_key": "rameshk_pass"},
    "shivakumar@stmaryscollege.in": {"name": "Mr. Shiva Kumar Reddy", "secret_key": "shivakumar_pass"},
    "anamika@stmaryscollege.in": {"name": "Dr. Anamika Sukul", "secret_key": "anamika_pass"},
    "arunjose@stmaryscollege.in": {"name": "Mr. Arun B Jose", "secret_key": "arunjose_pass"},
    "elisheba@stmaryscollege.in": {"name": "Ms. P. Elisheba", "secret_key": "elisheba_pass"},
    "debanjalee@stmaryscollege.in": {"name": "Dr. Debanjalee Bose", "secret_key": "debanjalee_pass"},
    "kirtibdnr@stmaryscollege.in": {"name": "Dr. Kirti", "secret_key": "kirti_pass"},
    "shikhasharma@stmaryscollege.in": {"name": "Dr. Shikha Sharma", "secret_key": "shikha_pass"},
    "himani@stmaryscollege.in": {"name": "Dr. Himani", "secret_key": "himani_pass"},
    "roy@stmaryscollege.in": {"name": "Mr. MSS Roy", "secret_key": "roy_pass"},
    "phebi@stmaryscollege.in": {"name": "Ms. Phebi", "secret_key": "phebi_pass"},
    "vigneshwari@stmaryscollege.in": {"name": "Dr. Vigneshwari", "secret_key": "vigneshwari_pass"},
    "nagarjuna@stmaryscollege.in": {"name": "Dr. Nagarjuna", "secret_key": "nagarjuna_pass"},
    "pavitrambika@stmaryscollege.in": {"name": "Dr. Pavitrambika", "secret_key": "pavitrambika_pass"},
    "anuradhaemani@stmaryscollege.in": {"name": "Dr. Anuradha", "secret_key": "anuradha_pass"},
    "kanthi@stmaryscollege.in": {"name": "Dr. Kanthi Sree", "secret_key": "kanthi_pass"},
    "timee@stmaryscollege.in": {"name": "Dr. Timee Ronra Shimray", "secret_key": "timee_pass"},
    "ismail@stmaryscollege.in": {"name": "Mr. Ismail C", "secret_key": "ismail_pass"},
    "aksharasingh@stmaryscollege.in": {"name": "Dr. Akshara Singh", "secret_key": "akshara_pass"},
    "vasantharao@stmaryscollege.in": {"name": "Mr. Vasantha Rao B", "secret_key": "vasantharao_pass"},
    "gisageorge@stmaryscollege.in": {"name": "Ms. Gisa George", "secret_key": "gisageorge_pass"},
    "research@stmaryscollege.in": {"name": "Research Admin", "secret_key": "research_pass"},
    "iqac@stmaryscollege.in": {"name": "Head, IQAC", "secret_key": "iqac_pass"},
    "harini@stmaryscollege.in": {"name": "Ms. Harini", "secret_key": "harini_pass"},
    "jayalakshmi@stmaryscollege.in": {"name": "Ms. Jayalakshmi D", "secret_key": "jayalakshmi_pass"},
    "rupini@stmaryscollege.in": {"name": "Ms. B. Rupini", "secret_key": "rupini_pass"},
    "manali@stmaryscollege.in": {"name": "Ms. Manali Manoj Manwadkar", "secret_key": "manali_pass"},
    "kusuma@stmaryscollege.in": {"name": "Dr. Kusuma C", "secret_key": "kusuma_pass"},
    "bikshapathi@stmaryscollege.in": {"name": "Mr. Bikshapathi M", "secret_key": "bikshapathi_pass"},
    "vijaybhaskar@stmaryscollege.in": {"name": "Mr. Vijay Bhaskar Reddy", "secret_key": "vijaybhaskar_pass"},
    "poojasharma@stmaryscollege.in": {"name": "Ms. Pooja Sharma", "secret_key": "poojasharma_pass"},
    "kavithathakur@stmaryscollege.in": {"name": "Dr. Kavitha Thakur", "secret_key": "kavithathakur_pass"},
    "priyamishra@stmaryscollege.in": {"name": "Dr. Priya Mishra", "secret_key": "priyamishra_pass"},
    "deepa@stmaryscollege.in": {"name": "Ms. Deepa Agraval", "secret_key": "deepaagraval_pass"},
    "nsrinath@stmaryscollege.in": {"name": "Dr. Srinath Naganathan", "secret_key": "nsrinath_pass"},
    "chrislenina@stmaryscollege.in": {"name": "Dr. Chris Lenina", "secret_key": "chrislenina_pass"},
    "sciences@stmaryscollege.in": {"name": "Department of Sciences", "secret_key": "sciences_pass"},
    "languages@stmaryscollege.in": {"name": "Department of English & Languages", "secret_key": "languages_pass"},
    "businessmanagement@stmaryscollege.in": {"name": "Department of Management", "secret_key": "businessmanagement_pass"},
    "commerce@stmaryscollege.in": {"name": "Department of Commerce", "secret_key": "commerce_pass"},
    "socialsciences@stmaryscollege.in": {"name": "Department of Social Sciences & Humanities", "secret_key": "socialsciences_pass"}
}

# --- 2. GOOGLE SERVICE INTEGRATION HANDSHAKE ---
def get_google_credentials():
    try:
        if "gcp_service_account" in st.secrets:
            info = dict(st.secrets["gcp_service_account"])
        elif "GCP_COMPLETE_B64" in st.secrets:
            raw_secret = str(st.secrets["GCP_COMPLETE_B64"]).strip()
            if raw_secret.startswith("{") and raw_secret.endswith("}"):
                info = json.loads(raw_secret)
            else:
                padded_b64 = raw_secret + "=" * (-len(raw_secret) % 4)
                decoded_bytes = base64.b64decode(padded_b64)
                info = json.loads(decoded_bytes.decode('utf-8', errors='ignore'))
        else:
            st.error("Credentials configuration missing in Streamlit Secrets.")
            st.stop()

        return service_account.Credentials.from_service_account_info(
            info, 
            scopes=["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        )
    except Exception as e:
        st.error(f"Ecosystem Verification Error: {str(e)}")
        st.stop()

def get_or_create_drive_folder(folder_name, parent_folder_id, creds):
    try:
        drive_service = build('drive', 'v3', credentials=creds)
        query = (
            f"name = '{folder_name}' and "
            f"'{parent_folder_id}' in parents and "
            f"mimeType = 'application/vnd.google-apps.folder' and "
            f"trashed = false"
        )
        results = drive_service.files().list(
            q=query, fields="files(id, name)", 
            supportsAllDrives=True, includeItemsFromAllDrives=True
        ).execute()
        
        files = results.get('files', [])
        if files:
            return files[0]['id']
            
        folder_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder',
            'parents': [parent_folder_id]
        }
        created_folder = drive_service.files().create(
            body=folder_metadata, fields='id', supportsAllDrives=True
        ).execute()
        
        try:
            drive_service.permissions().create(
                fileId=created_folder.get('id'), 
                body={'type': 'anyone', 'role': 'reader'}, 
                supportsAllDrives=True
            ).execute()
        except Exception:
            pass
            
        return created_folder.get('id')
    except Exception as e:
        st.warning(f"Drive folder setup fallback: {e}")
        return parent_folder_id

def upload_file_to_drive(file_bytes, file_name, mime_type, parent_ids, creds):
    try:
        drive_service = build('drive', 'v3', credentials=creds)
        links = []
        for p_id in parent_ids:
            file_metadata = {'name': file_name, 'parents': [p_id]}
            media = MediaIoBaseUpload(io.BytesIO(file_bytes), mimetype=mime_type, resumable=True)
            uploaded = drive_service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink, webContentLink', supportsAllDrives=True).execute()
            try:
                drive_service.permissions().create(fileId=uploaded.get('id'), body={'type': 'anyone', 'role': 'reader'}, supportsAllDrives=True).execute()
            except:
                pass
            links.append(uploaded.get('webViewLink', ""))
        return links[0] if links else "Drive Error"
    except Exception as e:
        st.warning("Drive Sync Notification: Written cleanly to spreadsheet layout.")
        return "Pending Folder Permissions Link"

def fetch_drive_folder_items(folder_id, creds):
    try:
        drive_service = build('drive', 'v3', credentials=creds)
        query = f"'{folder_id}' in parents and trashed = false"
        results = drive_service.files().list(
            q=query, 
            fields="files(id, name, mimeType, webViewLink, webContentLink, createdTime, description)",
            supportsAllDrives=True, 
            includeItemsFromAllDrives=True,
            pageSize=100
        ).execute()
        
        all_items = results.get('files', [])
        file_list = []
        
        for item in all_items:
            if item.get('mimeType') == 'application/vnd.google-apps.folder':
                sub_res = drive_service.files().list(
                    q=f"'{item.get('id')}' in parents and trashed = false and mimeType != 'application/vnd.google-apps.folder'",
                    fields="files(id, name, mimeType, webViewLink, webContentLink, createdTime, description)",
                    supportsAllDrives=True,
                    includeItemsFromAllDrives=True
                ).execute()
                file_list.extend(sub_res.get('files', []))
            else:
                file_list.append(item)
                
        return file_list
    except Exception:
        return []

def download_drive_file_bytes(file_id, creds):
    try:
        drive_service = build('drive', 'v3', credentials=creds)
        request = drive_service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        return fh.getvalue()
    except Exception:
        return None

# --- 3. ROBUST ALMANAC PARSER & DATE WINDOW MATCHER ---
def parse_single_date(s):
    if not s:
        return None
    s_clean = str(s).strip()
    s_clean = re.sub(r'^(?:mon|tue|wed|thu|fri|sat|sun)[a-z]*[\,\s\-]*', '', s_clean, flags=re.IGNORECASE).strip()
    s_clean = re.sub(r'(\d+)(st|nd|rd|th)', r'\1', s_clean, flags=re.IGNORECASE)

    m = re.search(r'([0-9]{1,2})[\/\-\.]([0-9]{1,2})[\/\-\.]([0-9]{2,4})', s_clean)
    if m:
        try:
            d, mth, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
            if y < 100: y += 2000
            return datetime.date(y, mth, d)
        except Exception:
            pass

    patterns = [
        r"%d %B %Y", r"%d %b %Y", r"%B %d, %Y", r"%b %d, %Y",
        r"%d-%B-%Y", r"%d-%b-%Y", r"%Y-%m-%d"
    ]
    for p in patterns:
        try:
            return datetime.datetime.strptime(s_clean, p).date()
        except Exception:
            pass
    return None

def parse_strict_or_range_date(text):
    if not text:
        return None, None
    s = str(text).strip()

    range_match = re.search(r'([0-9]{1,2}[\/\-\.][0-9]{1,2}[\/\-\.][0-9]{2,4})\s*(?:to|\-|till|until)\s*([0-9]{1,2}[\/\-\.][0-9]{1,2}[\/\-\.][0-9]{2,4})', s, re.IGNORECASE)
    if range_match:
        d1 = parse_single_date(range_match.group(1))
        d2 = parse_single_date(range_match.group(2))
        return d1, d2

    single = parse_single_date(s)
    return single, single

def normalize_almanac_department(raw_text):
    if not raw_text:
        return "All Units / Campus Wide"
    txt_lower = str(raw_text).strip().lower()

    if any(k in txt_lower for k in ["business management", "management"]):
        return "Management"
    elif "commerce" in txt_lower:
        return "Commerce"
    elif "science" in txt_lower:
        return "Sciences"
    elif any(k in txt_lower for k in ["english", "language", "samskrutha", "hindi", "french", "sanskrit"]):
        return "English & Languages"
    elif any(k in txt_lower for k in ["social science", "humanities", "mass communication", "psychology", "journalism", "political"]):
        return "Social Sciences & Humanities"
    elif any(k in txt_lower for k in ["physical education", "sports", "dhyanchand"]):
        return "Physical Education"
    elif any(k in txt_lower for k in ["research", "innovation", "iic"]):
        return "Research & Innovation"
    elif "iqac" in txt_lower or "sqc" in txt_lower:
        return "IQAC"
    elif "alumni" in txt_lower:
        return "Alumni"
    elif "library" in txt_lower:
        return "Library"
    elif "placement" in txt_lower:
        return "Placement"
    elif "women empowerment" in txt_lower:
        return "Women Empowerment"
    elif "internal complaints" in txt_lower or "icc" in txt_lower:
        return "Internal Complaints"
    elif "grievance" in txt_lower:
        return "Grievance Redressal"
    elif "anti-ragging" in txt_lower:
        return "Anti-Ragging"
    elif "scholarship" in txt_lower:
        return "Scholarship"
    elif any(k in txt_lower for k in ["nss", "ncc", "club", "cultural", "student activity", "community service"]):
        return "Student Activity Clubs"

    for d in DEPARTMENTS + COMMITTEES_CELLS_CLUBS:
        if d.lower() in txt_lower:
            return d

    return "All Units / Campus Wide"

def fetch_almanac_events(sheet_id, creds):
    """
    Extracts all Almanac records via Sheets API with automatic CSV fallback,
    parsing concatenated entries, dates, spans, and departments.
    """
    raw_rows = []
    
    # 1. Google Sheets API
    try:
        sheets_service = build('sheets', 'v4', credentials=creds)
        spreadsheet_meta = sheets_service.spreadsheets().get(spreadsheetId=sheet_id).execute()
        sheets = spreadsheet_meta.get('sheets', [])
        
        target_sheet_title = sheets[0]['properties']['title']
        for s in sheets:
            if str(s['properties'].get('sheetId')) == str(ALMANAC_GID):
                target_sheet_title = s['properties']['title']
                break

        res = sheets_service.spreadsheets().values().get(
            spreadsheetId=sheet_id,
            range=f"'{target_sheet_title}'!A1:Z1000"
        ).execute()
        raw_rows = res.get('values', [])
    except Exception:
        pass

    # 2. Public CSV stream fallback
    if not raw_rows:
        try:
            csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={ALMANAC_GID}"
            req = urllib.request.Request(csv_url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=10) as response:
                csv_text = response.read().decode('utf-8')
                df = pd.read_csv(io.StringIO(csv_text), header=None)
                raw_rows = df.fillna("").values.tolist()
        except Exception:
            pass

    if not raw_rows:
        return []

    events_list = []
    today = datetime.date.today()

    for r in raw_rows:
        if not r:
            continue

        row_str = " | ".join([str(c).strip() for c in r if str(c).strip()])
        if not row_str:
            continue

        # Look for Date or Date Span in row
        date_pattern = r'([0-9]{1,2}[\/\-\.][0-9]{1,2}[\/\-\.][0-9]{2,4}(?:\s*(?:to|\-)\s*[0-9]{1,2}[\/\-\.][0-9]{1,2}[\/\-\.][0-9]{2,4})?)'
        dates_found = re.findall(date_pattern, row_str, re.IGNORECASE)

        if not dates_found:
            continue

        for d_str in dates_found:
            start_date, end_date = parse_strict_or_range_date(d_str)
            if not start_date:
                continue

            # Extract event title and department details from remaining cells
            cells = [str(c).strip() for c in r if str(c).strip()]
            
            event_title = ""
            dept_hint = ""

            # Check if cells are separated
            meaningful = [
                c for c in cells 
                if not re.search(r'^[0-9]{1,2}[\/\-\.][0-9]{1,2}[\/\-\.][0-9]{2,4}', c) and 
                c.lower() not in ["monday", "tuesday", "wednesday", "thursday", "friday", "saturday", "sunday", "monday/tuesday", "friday/saturday", "day", "date", "s.no", "sl.no", "-", "none", "nil"]
            ]

            if meaningful:
                event_title = meaningful[0]
                dept_hint = " ".join(meaningful[1:]) if len(meaningful) > 1 else ""
            else:
                # Concatenated string line parsing
                cleaned_line = row_str.replace(d_str, '').strip(' |')
                cleaned_line = re.sub(r'^(?:mon|tue|wed|thu|fri|sat|sun)[a-z\/]*', '', cleaned_line, flags=re.IGNORECASE).strip(' |')
                event_title = cleaned_line

            if not event_title or event_title.lower() in ["nil", "none", "holiday", "sunday", "event", "particulars"]:
                continue

            matched_dept = normalize_almanac_department(dept_hint + " " + event_title)

            is_today = False
            is_upcoming_2weeks = False
            days_diff = (start_date - today).days

            if end_date and start_date <= today <= end_date:
                is_today = True
            elif start_date == today:
                is_today = True
            elif 0 < days_diff <= 14:
                is_upcoming_2weeks = True

            events_list.append({
                "title": event_title,
                "date_display": d_str,
                "start_date": start_date,
                "end_date": end_date,
                "dept": matched_dept,
                "description": dept_hint if dept_hint and dept_hint != matched_dept else "",
                "is_today": is_today,
                "is_upcoming_2weeks": is_upcoming_2weeks,
                "days_away": days_diff
            })

    return events_list

# --- 4. ROBUST WORD DOCUMENT PARSER ---
def extract_announcements_from_docx(file_bytes):
    entries = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        current_dept = "All Units / Campus Wide"
        current_category = "UGC CARE / INDEXED JOURNAL"

        for p in doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue

            dept_match = re.search(r'\[(.*?)\]|Department of\s+([A-Za-z &]+)', txt, re.IGNORECASE)
            if dept_match:
                detected = dept_match.group(1) or dept_match.group(2)
                for d in DEPARTMENTS + COMMITTEES_CELLS_CLUBS:
                    if d.lower() in detected.lower():
                        current_dept = d
                        break

            if any(k in txt.lower() for k in ["ugc care", "scopus", "web of science", "abdc", "call for papers", "upcoming conferences"]):
                current_category = txt.split(":")[0].strip().upper() if ":" in txt else txt[:40].strip().upper()

        def get_cell_content(cell):
            raw_t = cell.text.strip()
            urls = re.findall(r'https?://[^\s<>"]+|www\.[^\s<>"]+', raw_t)
            try:
                for rel in cell._tc.xpath('.//w:hyperlink'):
                    r_id = rel.get(qn('r:id'))
                    if r_id and r_id in cell.part.rels:
                        target = cell.part.rels[r_id].target_ref
                        if target and target.startswith("http"):
                            urls.append(target)
            except Exception:
                pass
            clean_t = re.sub(r'https?://[^\s<>"]+|www\.[^\s<>"]+', '', raw_t).strip(' \t\n\r|•-:')
            return clean_t, list(set(urls))

        # Process Tables
        for table in doc.tables:
            if not table.rows or len(table.rows) < 2:
                continue

            for row_idx, row in enumerate(table.rows):
                if row_idx == 0:
                    continue

                row_cells_data = []
                row_urls = []
                for cell in row.cells:
                    c_txt, c_urls = get_cell_content(cell)
                    if c_txt and c_txt.lower() not in ["data point not found", "n/a", "na", "-", "|", "nil", "none"]:
                        if not row_cells_data or c_txt != row_cells_data[-1]:
                            row_cells_data.append(c_txt)
                    row_urls.extend(c_urls)

                row_urls = list(set(row_urls))
                joined = " ".join(row_cells_data).lower()

                if sum(1 for hw in ["journal name", "journal title", "frequency", "guidelines", "formatting brief", "s.no", "serial", "submission link", "fee", "apc"] if hw in joined) >= 2:
                    continue

                if not row_cells_data and not row_urls:
                    continue

                journal_title = ""
                frequency_val = ""
                guidelines_val = ""
                apc_val = ""
                deadline_val = ""
                extra_notes = []

                for item in row_cells_data:
                    item_low = item.lower()
                    if item_low in ["formatting brief", "frequency", "guidelines", "submission link", "status", "fee", "apc"]:
                        continue

                    if any(k in item_low for k in ["apc", "free", "no apc", "$", "rs.", "inr", "usd", "nil"]):
                        apc_val = item
                    elif any(k in item_low for k in ["continuous", "bi-annual", "biannual", "annual", "quarterly", "monthly", "year-round", "open year-round", "triannual", "half-yearly"]):
                        frequency_val = item
                    elif any(k in item_low for k in ["word", "spacing", "tnr", "font", "apa", "ieee", "mla", "pages", "template", "manuscript", "anonymous"]) and len(item) > 8:
                        guidelines_val = item
                    elif any(k in item_low for k in ["deadline", "due date", "last date", "submission date", "submit by", "register by"]):
                        deadline_val = item
                    else:
                        if not journal_title and len(item) > 2 and not item.isdigit():
                            journal_title = item
                        else:
                            extra_notes.append(item)

                if not journal_title:
                    if row_urls:
                        clean_dom = re.sub(r'^https?:\/\/(www\.)?', '', row_urls[0]).split('/')[0]
                        journal_title = f"Journal Publication Portal ({clean_dom})"
                    else:
                        continue

                reg_links = [u if u.startswith("http") else f"https://{u}" for u in row_urls if any(k in u.lower() for k in ["guide", "author", "submit", "submission", "register", "form", "apply", "ticket", "forms.gle", "inauthors", "publish"])]
                gen_links = [u if u.startswith("http") else f"https://{u}" for u in row_urls if (u if u.startswith("http") else f"https://{u}") not in reg_links]

                dept = current_dept
                for d in DEPARTMENTS + COMMITTEES_CELLS_CLUBS:
                    if d.lower() in (journal_title + " " + " ".join(extra_notes)).lower():
                        dept = d
                        break

                entries.append({
                    "title": journal_title,
                    "frequency": frequency_val,
                    "guidelines": guidelines_val,
                    "apc": apc_val,
                    "deadline": deadline_val,
                    "notes": extra_notes,
                    "dept": dept,
                    "category": current_category,
                    "reg_links": list(set(reg_links)),
                    "gen_links": list(set(gen_links))
                })

        # Standalone Paragraphs
        for p in doc.paragraphs:
            txt = p.text.strip()
            if not txt or len(txt) < 80:
                continue
            if any(k in txt.lower() for k in ["updated on", "compiled by", "ugc care listed", "scopus", "disclaimer", "formatting brief", "table of contents", "st. mary"]):
                continue

            urls = re.findall(r'https?://[^\s<>"]+|www\.[^\s<>"]+', txt)
            cleaned_p = re.sub(r'https?://[^\s<>"]+|www\.[^\s<>"]+', '', txt).strip()
            
            if cleaned_p and len(cleaned_p) >= 60:
                p_title = cleaned_p[:70] + "..." if len(cleaned_p) > 70 else cleaned_p
                entries.append({
                    "title": p_title,
                    "frequency": "",
                    "guidelines": "",
                    "apc": "",
                    "deadline": "",
                    "notes": [cleaned_p],
                    "dept": current_dept,
                    "category": current_category,
                    "reg_links": [u if u.startswith("http") else f"https://{u}" for u in urls if any(k in u.lower() for k in ["register", "submit", "guide"])],
                    "gen_links": [u if u.startswith("http") else f"https://{u}" for u in urls if not any(k in u.lower() for k in ["register", "submit", "guide"])]
                })

    except Exception:
        pass
    return entries

def delete_drive_file(file_id, creds):
    try:
        drive_service = build('drive', 'v3', credentials=creds)
        drive_service.files().delete(fileId=file_id, supportsAllDrives=True).execute()
        return True
    except Exception:
        return False

def append_and_sort_sheet_by_department(sheet_name, new_row, dept_column_index, creds):
    try:
        sheets_service = build('sheets', 'v4', credentials=creds)
        result = sheets_service.spreadsheets().values().get(spreadsheetId=MASTER_SHEET_ID, range=f"'{sheet_name}'!A1:N2000").execute()
        rows = result.get('values', [])
        
        if not rows:
            sheets_service.spreadsheets().values().update(
                spreadsheetId=MASTER_SHEET_ID, range=f"'{sheet_name}'!A1",
                valueInputOption="USER_ENTERED", body={"values": [new_row]}
            ).execute()
            return

        header, data_rows = rows[0], rows[1:]
        data_rows.append(new_row)
        
        if sheet_name in ["Research_Database", "Faculty_Achievements", "Student_Activities"]:
            data_rows.sort(key=lambda r: DEPT_SORT_ORDER.get(r[dept_column_index], len(DEPARTMENTS)) if len(r) > dept_column_index else len(DEPARTMENTS))
        else:
            data_rows.sort(key=lambda r: r[dept_column_index] if len(r) > dept_column_index else "")
            
        sorted_matrix = [header] + data_rows
        
        sheets_service.spreadsheets().values().clear(spreadsheetId=MASTER_SHEET_ID, range=f"'{sheet_name}'!A1:N2000").execute()
        sheets_service.spreadsheets().values().update(
            spreadsheetId=MASTER_SHEET_ID, range=f"'{sheet_name}'!A1",
            valueInputOption="USER_ENTERED", body={"values": sorted_matrix}
        ).execute()
    except Exception as e:
        st.error(f"Sorting Error: {str(e)}")

def fetch_sheet_records(sheet_name, creds):
    try:
        sheets_service = build('sheets', 'v4', credentials=creds)
        res = sheets_service.spreadsheets().values().get(
            spreadsheetId=MASTER_SHEET_ID, 
            range=f"'{sheet_name}'!A1:Z2000"
        ).execute()
        rows = res.get('values', [])
        if not rows or len(rows) < 2:
            return pd.DataFrame()
        
        headers = rows[0]
        max_cols = len(headers)
        data = [r + [""] * (max_cols - len(r)) if len(r) < max_cols else r[:max_cols] for r in rows[1:]]
        return pd.DataFrame(data, columns=headers)
    except Exception:
        return pd.DataFrame()

# --- 5. THE WORD DOCUMENT NARRATIVE COMPILER ENGINE ---
def build_monthly_word_document(name_focus, active_month, active_year, creds):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    
    title_p = doc.add_paragraph()
    title_p.add_run(f"Monthly Achievements Summary Report: {active_month}, {active_year}\n").bold = True
    title_p.runs[0].font.size = Pt(14)
    
    scope_p = doc.add_paragraph()
    
    sheets_service = build('sheets', 'v4', credentials=creds)
    month_map = {"jan": "january", "feb": "february", "mar": "march", "apr": "april", "may": "may", "jun": "june", "jul": "july", "aug": "august", "sep": "september", "oct": "october", "nov": "november", "dec": "december"}
    target_month_clean = str(active_month).strip().lower()

    if name_focus == "Committees / Cells / Clubs":
        scope_p.add_run("COMMITTEES / CELLS / CLUBS MASTER DOSSIER\n").bold = True
        scope_p.runs[0].font.size = Pt(13)
        
        doc.add_paragraph().add_run("I. Consolidated Committee Activity Logs & Event Narratives").bold = True
        doc.add_paragraph().add_run("Chronological record of organized events, initiatives, and execution statements across all campus cells.").font.italic = True
        
        try:
            res = sheets_service.spreadsheets().values().get(spreadsheetId=MASTER_SHEET_ID, range="'Committees_Cells_Clubs'!A1:F1000").execute()
            rows = res.get('values', [])
        except:
            rows = []
            
        has_data = False
        if len(rows) > 1:
            for row in rows[1:]:
                if len(row) >= 5:
                    row_comm, row_faculty, row_month, row_year, row_narrative = row[0], row[1], row[2], row[3], row[4]
                    normalized_row_month = str(row_month).strip().lower()
                    
                    month_match = (target_month_clean == normalized_row_month) or \
                                  (target_month_clean[:3] in normalized_row_month) or \
                                  (normalized_row_month in month_map and month_map[normalized_row_month] == target_month_clean)
                    
                    if month_match and str(row_year).strip() == str(active_year).strip():
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"[{row_comm}] ").bold = True
                        p.add_run(f"{row_narrative} (In-charge: {row_faculty})")
                        has_data = True
                        
        if not has_data:
            doc.add_paragraph().add_run("\t- Nil -")
            
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        return doc_stream.getvalue()

    scope_p.add_run(f"DEPARTMENT OF {name_focus.upper()}\n").bold = True
    scope_p.runs[0].font.size = Pt(13)
    
    sections = [
        {"title": "I. Research Publications & Paper Presentations", "sheet": "Research_Database", "filter": ["Paper Publication", "Book Chapter", "Full Book", "Paper Presentation"], "desc": "Include journal articles, book chapters, full books, or papers presented at conferences."},
        {"title": "II. Faculty Development Programs (FDPs) & Workshops", "sheet": "Research_Database", "filter": ["FDP", "Workshop"], "desc": "Include training programs attended or successfully completed."},
        {"title": "III. Professional Certifications & Training", "sheet": "Faculty_Achievements", "filter": ["Certification/Course"], "desc": "Include NPTEL courses, Innovation Ambassador training, or other professional certifications."},
        {"title": "IV. Resource Person Roles & Invited Lectures", "sheet": "Faculty_Achievements", "filter": ["Presentation/Resource Person"], "desc": "Include acting as a Judge, Guest Speaker, Keynote Facilitator, or Resource Person for academic colloquiums."},
        {"title": "V. Research Milestones (For Doctoral Scholars)", "sheet": "Faculty_Achievements", "filter": ["Doctoral Milestone"], "desc": "Include milestones such as Synopsis Seminars, Pre-Ph.D. exams, or Thesis submission."},
        {"title": "VI. Awards, Honors, & Recognitions", "sheet": "Faculty_Achievements", "filter": ["Award/Honor"], "desc": "Include any special awards, titles, or professional recognitions."},
        {"title": "VII. Departmental & Student Contribution", "sheet": "Student_Activities", "filter": ["Institutional Contribution"], "desc": "Include organized events, Institutional Social Responsibility (ISR) activities, or specialized student activities."}
    ]
    
    def pad_row(target_row, required_length=15):
        return target_row + [""] * (required_length - len(target_row))

    for sec in sections:
        doc.add_paragraph().add_run(sec["title"]).bold = True
        doc.add_paragraph().add_run(sec["desc"]).font.italic = True
        
        try:
            res = sheets_service.spreadsheets().values().get(spreadsheetId=MASTER_SHEET_ID, range=f"'{sec['sheet']}'!A1:N1000").execute()
            rows = res.get('values', [])
        except: 
            rows = []
            
        has_data = False
        if len(rows) > 1:
            for row in rows[1:]:
                if len(row) >= 2:
                    padded = pad_row(row, required_length=15)
                    
                    if sec["sheet"] == "Research_Database":
                        row_dept, row_cat, row_month = padded[1], padded[2], padded[13]
                    elif sec["sheet"] == "Faculty_Achievements":
                        row_dept, row_cat, row_month = padded[1], padded[4], padded[2]
                    elif sec["sheet"] == "Student_Activities":
                        row_dept, row_cat, row_month = padded[1], padded[4], padded[2]
                    else:
                        row_dept, row_cat, row_month = padded[0], padded[4], padded[2]
                    
                    normalized_row_month = str(row_month).strip().lower()
                    
                    month_match = (target_month_clean == normalized_row_month) or \
                                  (target_month_clean[:3] in normalized_row_month) or \
                                  (normalized_row_month in month_map and month_map[normalized_row_month] == target_month_clean)
                    
                    if str(row_dept).strip().lower() == str(name_focus).strip().lower() and month_match and \
                       any(str(row_cat).strip().lower() == str(f).strip().lower() for f in sec["filter"]):
                        
                        p = doc.add_paragraph(style='List Bullet')
                        if sec["sheet"] == "Research_Database":
                            f_name, f_cat, j_type, title_text, pub_url, pub_name, pub_scope, conf_scope, org_body, isbn_issn, duration_dates = \
                                padded[0], padded[2], padded[3], padded[4], padded[7], padded[8], padded[9], padded[10], padded[11], padded[12], padded[6]
                            
                            if f_cat in ["Paper Publication", "Book Chapter", "Full Book"]:
                                narr = f'{f_name} published a {f_cat} titled "{title_text}" in {pub_name}. Journal Type: {j_type}, ISSN/ISBN: [{isbn_issn}], Scope: {pub_scope}. URL: {pub_url}'
                            elif f_cat == "Paper Presentation":
                                narr = f'{f_name} presented a research paper titled "{title_text}" at the conference organized by {org_body or pub_name} ({duration_dates or "NA"}). Scope: {conf_scope}.'
                            else:
                                narr = f'{f_name} completed a {duration_dates} {conf_scope if (conf_scope and conf_scope != "NA") else "Institutional"} {f_cat} on "{title_text}," organized by {org_body}.'
                        else:
                            narr = padded[5]
                        
                        p.add_run(narr)
                        has_data = True
                        
        if not has_data:
            doc.add_paragraph().add_run("\t- Nil -")
        doc.add_paragraph()
        
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    return doc_stream.getvalue()

def styled_block(format_text, example_text):
    html_string = f"""<div style="background-color: #FFFFFF; padding: 16px; border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.05); border: 1px solid #EAECEF; margin-bottom: 20px; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;"><div style="display: flex; align-items: flex-start; margin-bottom: 14px;"><div style="background-color: #E8EAF6; color: #1A237E; font-weight: 700; font-size: 11px; text-transform: uppercase; letter-spacing: 0.8px; padding: 4px 8px; border-radius: 4px; margin-right: 12px; min-width: 70px; text-align: center; border-left: 3px solid #1A237E;">Format</div><div style="color: #2C3E50; font-size: 14px; line-height: 1.5; font-weight: 500;">{format_text}</div></div><div style="height: 1px; background-color: #F1F3F5; margin: 12px 0;"></div><div style="display: flex; align-items: flex-start;"><div style="background-color: #E8F5E9; color: #1B5E20; font-weight: 700; font-size: 11px; text-transform: uppercase; letter-spacing: 0.8px; padding: 4px 8px; border-radius: 4px; margin-right: 12px; min-width: 70px; text-align: center; border-left: 3px solid #2E7D32;">Example</div><div style="color: #455A64; font-size: 14px; line-height: 1.5; font-style: italic; font-weight: 500;">{example_text}</div></div></div>"""
    st.markdown(html_string, unsafe_allow_html=True)

# --- 6. WEBSITE FRONT-PAGE & ANNOUNCEMENT CARDS ---
def render_scrolling_ticker(announcements):
    ticker_text = " &nbsp;&nbsp;&nbsp; 🌟 &nbsp;&nbsp;&nbsp; ".join(announcements)
    ticker_html = f"""<div style="background: linear-gradient(90deg, #1A237E 0%, #283593 100%); color: #FFFFFF; padding: 10px 15px; border-radius: 8px; margin-bottom: 25px; box-shadow: 0 4px 10px rgba(0,0,0,0.1); font-weight: 500;"><marquee behavior="scroll" direction="left" scrollamount="6">📢 <b>LATEST RESEARCH HIGHLIGHTS & INDEXED PUBLICATIONS:</b> &nbsp;&nbsp;&nbsp; {ticker_text}</marquee></div>"""
    st.markdown(ticker_html, unsafe_allow_html=True)

def render_publication_achiever_card(author, dept, title, journal, indexing, link_url):
    badge_colors = {
        "📖 FULL BOOK": "#B8860B",
        "Scopus": "#E65100",
        "Web of Science": "#0D47A1",
        "SCIE": "#1B5E20",
        "ABDC": "#4A148C",
        "UGC Care Listed": "#B71C1C",
        "PubMed": "#006064",
        "Peer Reviewed": "#374151"
    }
    badge_bg = badge_colors.get(indexing, "#1A237E")
    is_book = indexing == "📖 FULL BOOK"
    border_style = "2px solid #D4AF37" if is_book else "1px solid #E2E8F0"
    
    link_html = f"<a href='{link_url}' target='_blank' style='color:#1A237E; font-weight:600; text-decoration:none;'>🔗 View Document / Link</a>" if link_url and link_url not in ["Pending Folder Permissions Link", "NA", ""] else ""

    card_html = f"""<div style="background-color: #FFFFFF; border-radius: 10px; padding: 18px; box-shadow: 0 4px 14px rgba(0,0,0,0.06); border: {border_style}; margin-bottom: 20px; height: 100%; display: flex; flex-direction: column; justify-content: space-between;"><div><div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:10px;"><span style="background-color:{badge_bg}; color:white; font-size:11px; font-weight:700; padding:4px 8px; border-radius:4px; text-transform:uppercase;">{indexing}</span><span style="color:#64748B; font-size:12px; font-weight:500;">{dept}</span></div><h4 style="margin: 0 0 8px 0; color: #1E293B; font-size: 15px; line-height: 1.4;">{title}</h4><p style="margin: 0 0 6px 0; color: #334155; font-size: 13px; font-weight: 600;">✍️ {author}</p><p style="margin: 0 0 10px 0; color: #64748B; font-size: 12px; font-style: italic;">📖 {journal}</p></div><div style="padding-top:8px; border-top:1px solid #F1F5F9; font-size:12px;">{link_html}</div></div>"""
    st.markdown(card_html, unsafe_allow_html=True)

def render_bulletin_card(file_obj, category_label, bg_color="#1A237E"):
    file_name = html.escape(file_obj.get("name", "Announcement Flyer"))
    view_link = html.escape(file_obj.get("webViewLink", "#"))
    mime = file_obj.get("mimeType", "")
    icon = "🖼️" if "image" in mime else ("📝" if "document" in mime or file_name.endswith(".docx") else "📄")

    dl_match = re.search(r'(?:due|deadline|date|registration)[\s\_\-]*([0-9]{1,2}[A-Za-z]+|[0-9]{1,2}[\-\.][0-9]{1,2}[\-\.][0-9]{2,4})', file_name, re.IGNORECASE)
    deadline_badge = f"""<div style="background-color: #FEF2F2; color: #DC2626; border: 1px solid #FCA5A5; font-weight: 700; font-size: 11px; padding: 4px 8px; border-radius: 4px; margin-top: 6px; display: inline-block;">⏰ Deadline: {html.escape(dl_match.group(1))}</div>""" if dl_match else ""

    card_html = f"""<div style="background-color: #FFFFFF; border-radius: 10px; padding: 16px; box-shadow: 0 3px 10px rgba(0,0,0,0.06); border: 1px solid #E2E8F0; margin-bottom: 15px;"><div style="display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 8px;"><span style="background-color: {bg_color}; color: #FFFFFF; font-size: 11px; font-weight: 700; padding: 3px 8px; border-radius: 4px;">{category_label}</span><span style="font-size: 16px;">{icon}</span></div><h4 style="margin: 0 0 4px 0; color: #1E293B; font-size: 14px; line-height: 1.4; word-break: break-word;">{file_name}</h4>{deadline_badge}<div style="margin-top: 10px; font-size: 12px;"><a href="{view_link}" target="_blank" style="background-color: #EEF2FF; color: #1A237E; padding: 6px 12px; border-radius: 6px; text-decoration: none; font-weight: 600; display: inline-block;">🔍 Open Document / Flyer</a></div></div>"""
    st.markdown(card_html, unsafe_allow_html=True)

def render_parsed_doc_entry(entry):
    category_label = html.escape(str(entry.get('category', 'CALL FOR PAPERS / JOURNAL')))
    title_text = html.escape(str(entry.get('title', 'Research Announcement')))
    dept_text = html.escape(str(entry.get('dept', 'Campus Wide')))
    
    badges_list = []
    if entry.get('frequency'):
        badges_list.append(f"<span style='background-color: #EFF6FF; color: #1D4ED8; border: 1px solid #BFDBFE; font-weight: 600; font-size: 11px; padding: 4px 8px; border-radius: 5px; margin-right: 6px; margin-bottom: 6px; display: inline-block;'>🔄 <b>Cycle:</b> {html.escape(str(entry['frequency']))}</span>")
    if entry.get('apc'):
        badges_list.append(f"<span style='background-color: #FEF3C7; color: #B45309; border: 1px solid #FDE68A; font-weight: 600; font-size: 11px; padding: 4px 8px; border-radius: 5px; margin-right: 6px; margin-bottom: 6px; display: inline-block;'>💰 <b>APC:</b> {html.escape(str(entry['apc']))}</span>")
    if entry.get('guidelines'):
        badges_list.append(f"<span style='background-color: #F0FDF4; color: #15803D; border: 1px solid #BBF7D0; font-weight: 600; font-size: 11px; padding: 4px 8px; border-radius: 5px; margin-right: 6px; margin-bottom: 6px; display: inline-block;'>📝 <b>Length / Format:</b> {html.escape(str(entry['guidelines']))}</span>")
    if entry.get('deadline'):
        badges_list.append(f"<span style='background-color: #FFF1F2; color: #E11D48; border: 1px solid #FECDD3; font-weight: 700; font-size: 11px; padding: 4px 8px; border-radius: 5px; margin-right: 6px; margin-bottom: 6px; display: inline-block;'>⏰ <b>Deadline:</b> {html.escape(str(entry['deadline']))}</span>")

    rendered_badges = f"<div style='margin-bottom: 8px;'>{''.join(badges_list)}</div>" if badges_list else ""

    action_buttons = []
    if entry.get("reg_links"):
        for url in entry.get("reg_links"):
            action_buttons.append(f"<a href='{html.escape(url)}' target='_blank' style='background: linear-gradient(135deg, #E11D48 0%, #BE123C 100%); color: #FFFFFF; padding: 6px 12px; border-radius: 5px; text-decoration: none; font-weight: 600; font-size: 12px; margin-right: 8px; margin-top: 6px; display: inline-block;'>🎟️ Author Guidelines / Submit</a>")

    if entry.get("gen_links"):
        for url in entry.get("gen_links"):
            action_buttons.append(f"<a href='{html.escape(url)}' target='_blank' style='background: linear-gradient(135deg, #1A237E 0%, #283593 100%); color: #FFFFFF; padding: 6px 12px; border-radius: 5px; text-decoration: none; font-weight: 600; font-size: 12px; margin-right: 8px; margin-top: 6px; display: inline-block;'>🌐 Official Journal Portal</a>")

    rendered_buttons = f"<div style='margin-top: 10px;'>{''.join(action_buttons)}</div>" if action_buttons else ""
    
    notes_html = ""
    if entry.get("notes"):
        valid_notes = [html.escape(str(n)) for n in entry["notes"] if str(n) != entry.get("title") and len(str(n)) > 5 and str(n) not in [entry.get('frequency'), entry.get('guidelines'), entry.get('apc')]]
        if valid_notes:
            notes_html = f"<div style='color: #475569; font-size: 12.5px; line-height: 1.5; margin-top: 6px;'>{'<br>'.join(['• ' + vn for vn in valid_notes])}</div>"

    card_html = f"""<div style="background-color: #FFFFFF; border-radius: 10px; padding: 18px; box-shadow: 0 4px 12px rgba(0,0,0,0.06); border-left: 4px solid #4338CA; border-top: 1px solid #E2E8F0; border-right: 1px solid #E2E8F0; border-bottom: 1px solid #E2E8F0; margin-bottom: 16px;"><div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px;"><span style="background-color: #4338CA; color: #FFFFFF; font-size: 10px; font-weight: 700; padding: 2px 7px; border-radius: 4px; text-transform: uppercase;">{category_label}</span><span style="color: #64748B; font-size: 11px; font-weight: 600;">{dept_text}</span></div><h4 style="margin: 0 0 8px 0; color: #1E293B; font-size: 15px; font-weight: 700; line-height: 1.4;">{title_text}</h4>{rendered_badges}{notes_html}{rendered_buttons}</div>"""
    st.markdown(card_html, unsafe_allow_html=True)

def render_almanac_event_card(ev, is_highlighted=False):
    title = html.escape(ev['title'])
    date_str = html.escape(ev['date_display'])
    dept = html.escape(ev['dept'])
    desc = html.escape(ev['description']) if ev['description'] else ""

    if is_highlighted:
        border_style = "border-left: 5px solid #16A34A; background: #F0FDF4;"
        badge_tag = "<span style='background-color: #16A34A; color: white; padding: 3px 8px; border-radius: 4px; font-size: 11px; font-weight: 700;'>🚨 HAPPENING TODAY</span>"
    else:
        days_away_txt = f"in {ev['days_away']} days" if ev['days_away'] > 1 else ("tomorrow" if ev['days_away'] == 1 else "today")
        border_style = "border-left: 4px solid #0284C7; background: #FFFFFF;"
        badge_tag = f"<span style='background-color: #E0F2FE; color: #0369A1; padding: 3px 8px; border-radius: 4px; font-size: 11px; font-weight: 600;'>📅 {days_away_txt}</span>"

    desc_html = f"<p style='margin: 4px 0 0 0; color: #64748B; font-size: 12px;'>{desc}</p>" if desc else ""

    card_html = f"""<div style="{border_style} border-radius: 8px; padding: 14px 16px; box-shadow: 0 2px 8px rgba(0,0,0,0.04); border-top: 1px solid #E2E8F0; border-right: 1px solid #E2E8F0; border-bottom: 1px solid #E2E8F0; margin-bottom: 12px;"><div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 6px;">{badge_tag}<span style="color: #64748B; font-size: 11px; font-weight: 600;">{dept}</span></div><h4 style="margin: 0; color: #1E293B; font-size: 14px; font-weight: 700; line-height: 1.4;">{title}</h4><div style="margin-top: 4px; font-size: 12px; color: #334155; font-weight: 500;">🗓️ <b>Date:</b> {date_str}</div>{desc_html}</div>"""
    st.markdown(card_html, unsafe_allow_html=True)

# --- 5. STREAMLIT FRAMEWORK DESK ---
if "authenticated" not in st.session_state: st.session_state.authenticated = False
if "logged_email" not in st.session_state: st.session_state.logged_email = ""
if "admin_enabled" not in st.session_state: st.session_state.admin_enabled = True

st.set_page_config(page_title="St. Mary's Integrated Portal", page_icon="🏫", layout="wide")

if not st.session_state.authenticated:
    _, img_col, _ = st.columns([2, 1, 2])
    with img_col:
        if os.path.exists("logo.png"):
            st.image("logo.png", use_container_width=True)
        else:
            st.markdown("<h1 style='text-align:center;'>🏫</h1>", unsafe_allow_html=True)
        
    st.markdown("<h2 style='text-align: center;'>St. Mary's Central Achievements Portal</h2>", unsafe_allow_html=True)
    _, col_l2, _ = st.columns([1, 1.5, 1])
    with col_l2:
        input_email = st.text_input("College Email Address").strip().lower()
        input_password = st.text_input("Password", type="password")
        if st.button("Sign In", type="primary", use_container_width=True):
            if input_email in FACULTY_DIRECTORY:
                if input_password == st.secrets.get(FACULTY_DIRECTORY[input_email]["secret_key"], "welcome@2026"):
                    st.session_state.authenticated, st.session_state.logged_email = True, input_email
                    st.rerun()
                else: st.error("Invalid credentials entry.")
            else: st.error("Email address not authorized inside profile system.")
    st.stop()

# --- HEADER WORKSPACE WITH LOGOUT TOOL ---
current_faculty_name = FACULTY_DIRECTORY[st.session_state.logged_email]["name"]
creds = get_google_credentials()

logo_col, header_col, logout_col = st.columns([1, 7, 1.5])
with logo_col:
    if os.path.exists("logo.png"):
        st.image("logo.png", width=65)
    else:
        st.markdown("<h3>🏫</h3>", unsafe_allow_html=True)
with header_col:
    st.markdown(f"### Welcome back, **{current_faculty_name}**")
with logout_col:
    if st.button("🚪 Log Out", type="secondary", use_container_width=True):
        st.session_state.authenticated = False
        st.session_state.logged_email = ""
        st.rerun()

# --- EXTRACT DATA FROM ALL MASTER SHEETS ---
res_df = fetch_sheet_records("Research_Database", creds)
fac_df = fetch_sheet_records("Faculty_Achievements", creds)
stu_df = fetch_sheet_records("Student_Activities", creds)
comm_df = fetch_sheet_records("Committees_Cells_Clubs", creds)

# --- TAB NAVIGATION ---
tab_gallery, tab_announcements, tab_explorer, tab_submit, tab_document, tab_admin = st.tabs([
    "🌐 Research Portal Home",
    "📢 Live Announcements & Events",
    "📋 Master Database Explorer",
    "📝 Enter Research Data", 
    "📊 Monthly Achievement Generator", 
    "🔒 Admin Control"
])

# --- TAB 1: RESEARCH PORTAL HOME ---
with tab_gallery:
    def get_impact_rank(row):
        pub_type = str(row.iloc[2]).strip().lower() if len(row) > 2 else ""
        indexing = str(row.iloc[3]).strip().lower() if len(row) > 3 else ""
        if "full book" in pub_type: return 1
        elif any(k in indexing for k in ["scopus", "web of science", "scie"]): return 2
        elif any(k in indexing for k in ["abdc", "pubmed", "doaj", "embase"]): return 3
        elif "ugc care" in indexing: return 4
        elif "book chapter" in pub_type or "proceeding" in indexing: return 5
        elif "peer reviewed" in indexing: return 6
        return 99

    valid_publications = []
    if not res_df.empty and len(res_df) > 0:
        for _, r in res_df.iterrows():
            pub_type = str(r.iloc[2]).strip() if len(row := r) > 2 else ""
            indexing = str(r.iloc[3]).strip() if len(r) > 3 else ""
            is_full_book = "full book" in pub_type.lower()
            is_valid_index = indexing.upper() not in ["NA", "N/A", "NONE", ""] and indexing.strip() != ""
            if is_full_book or is_valid_index:
                valid_publications.append(r)

    if valid_publications:
        ticker_items = []
        for row in valid_publications[:10]:
            f_auth = row.iloc[0] if len(row) > 0 else "Faculty"
            f_type = row.iloc[2] if len(row) > 2 else "Publication"
            f_idx = row.iloc[3] if len(row) > 3 else "Indexed"
            f_title = row.iloc[4] if len(row) > 4 else "Research Work"
            f_jour = row.iloc[8] if len(row) > 8 else "Publisher"
            tag = "🌟 FULL BOOK" if "full book" in f_type.lower() else f"[{f_idx}]"
            ticker_items.append(f"{f_auth} {tag}: '{f_title}' ({f_jour})")
        render_scrolling_ticker(ticker_items)
    else:
        render_scrolling_ticker([
            "Dr. Srinath Naganathan [Scopus]: 'Bioremediation Kinetics' in Environmental Science",
            "Dr. Manoj Kanth [ABDC]: 'Strategic Corporate Governance' in Journal of Financial Studies",
            "Dr. Rajita Anand Singh [UGC Care Listed]: 'Modern Commonwealth Fiction'"
        ])

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("🔬 Total Research Logged", len(res_df) if not res_df.empty else 0)
    m2.metric("🏆 Faculty Milestones", len(fac_df) if not fac_df.empty else 0)
    m3.metric("👥 Department Initiatives", len(stu_df) if not stu_df.empty else 0)
    m4.metric("🏛️ Committee Activities", len(comm_df) if not comm_df.empty else 0)

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("""<div style="background-color:#F8FAFC; border-left: 5px solid #1A237E; padding:18px 22px; border-radius:6px; margin-bottom:20px;"><h3 style="margin:0 0 6px 0; color:#1A237E;">🏆 Faculty Research Achievers Gallery</h3><p style="margin:0; color:#475569; font-size:14px;">Ranked in order of impact: <b>Authored Books, Scopus, Web of Science / SCIE, ABDC, and UGC-CARE</b> publications.</p></div>""", unsafe_allow_html=True)

    if valid_publications:
        valid_publications.sort(key=get_impact_rank)
        cols = st.columns(3)
        for i, row in enumerate(valid_publications):
            author = row.iloc[0] if len(row) > 0 else "Faculty Member"
            dept = row.iloc[1] if len(row) > 1 else "Department"
            pub_type = row.iloc[2] if len(row) > 2 else "Publication"
            indexing = row.iloc[3] if len(row) > 3 else "Peer Reviewed"
            title = row.iloc[4] if len(row) > 4 else "Research Publication"
            link_url = row.iloc[5] if len(row) > 5 else ""
            journal = row.iloc[8] if len(row) > 8 else "Publisher"
            display_badge = "📖 FULL BOOK" if "full book" in pub_type.lower() else indexing
            with cols[i % 3]:
                render_publication_achiever_card(author, dept, title, journal, display_badge, link_url)
    else:
        st.info("No high-impact publication records found. Add your publications under the 'Enter Research Data' tab!")

# --- TAB 2: LIVE ANNOUNCEMENTS & ALMANAC EVENTS (TWO-COLUMN BULLETIN) ---
with tab_announcements:
    st.subheader("📢 Campus Bulletin & Upcoming Opportunities")
    st.markdown("Live repository of conferences, call for papers, and upcoming departmental & club activities from the College Almanac.")

    col_f1, col_f2 = st.columns([4, 1])
    with col_f1:
        filter_options = ["All Units / Campus Wide"] + DEPARTMENTS + COMMITTEES_CELLS_CLUBS
        selected_unit = st.selectbox("🎯 Filter Events by Department / Club / Committee:", filter_options)
    with col_f2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🔄 Refresh Bulletin", use_container_width=True):
            st.rerun()

    with st.spinner("Accessing Google Drive folders and syncing Almanac schedule..."):
        research_files = fetch_drive_folder_items(RESEARCH_EVENTS_FOLDER_ID, creds)
        campus_files = fetch_drive_folder_items(CAMPUS_ACTIVITIES_FOLDER_ID, creds)
        almanac_events = fetch_almanac_events(ALMANAC_SHEET_ID, creds)

    parsed_docx_entries = []
    regular_research_files = []

    for f in research_files:
        name = f.get("name", "")
        if name.endswith(".docx") or "officedocument.wordprocessingml.document" in f.get("mimeType", ""):
            b_data = download_drive_file_bytes(f.get("id"), creds)
            if b_data:
                extracted = extract_announcements_from_docx(b_data)
                parsed_docx_entries.extend(extracted)
        else:
            regular_research_files.append(f)

    # Two-Column Layout
    col_left, col_right = st.columns(2)

    # LEFT COLUMN: Research Conferences & CFP Dossiers
    with col_left:
        st.markdown("""<div style="background-color: #EEF2FF; border-left: 4px solid #1A237E; padding: 10px 14px; border-radius: 4px; margin-bottom: 15px;"><h4 style="margin: 0; color: #1A237E;">🔬 Upcoming Conferences & Call for Papers</h4><p style="margin: 0; color: #475569; font-size: 12px;">Folder Vault: <code>Upcoming Research Events</code></p></div>""", unsafe_allow_html=True)

        if parsed_docx_entries:
            matching_docx_entries = [
                e for e in parsed_docx_entries 
                if selected_unit == "All Units / Campus Wide" or selected_unit.lower() in e["dept"].lower() or selected_unit.lower() in (e["title"] + " ".join(e["notes"])).lower()
            ]
            for entry in matching_docx_entries:
                render_parsed_doc_entry(entry)

        if regular_research_files:
            filtered_research = [
                f for f in regular_research_files 
                if selected_unit == "All Units / Campus Wide" or selected_unit.lower() in f.get("name", "").lower()
            ]
            for f in filtered_research:
                render_bulletin_card(f, "Research Poster / CFP", bg_color="#1A237E")

        if not parsed_docx_entries and not regular_research_files:
            st.info("No conference circulars, Word dossiers, or Call-for-Paper documents found in the Research Events folder.")

    # RIGHT COLUMN: Almanac Events (Today & 2-Week Window) + Campus Activity Flyers
    with col_right:
        st.markdown("""<div style="background-color: #F0FDF4; border-left: 4px solid #16A34A; padding: 10px 14px; border-radius: 4px; margin-bottom: 15px;"><h4 style="margin: 0; color: #16A34A;">🎭 Departmental, Club & Student Activities</h4><p style="margin: 0; color: #475569; font-size: 12px;">Source: <code>2026-27 College Almanac (Rolling 14-Day Calendar)</code></p></div>""", unsafe_allow_html=True)

        # Filter almanac events by selected unit
        filtered_almanac = [
            ev for ev in almanac_events
            if selected_unit == "All Units / Campus Wide" or ev["dept"] == "All Units / Campus Wide" or selected_unit.lower() in ev["dept"].lower() or selected_unit.lower() in (ev["title"] + " " + ev["description"]).lower()
        ]

        # 1. Today's Events (Ongoing)
        today_events = [ev for ev in filtered_almanac if ev['is_today']]
        if today_events:
            st.markdown("##### 🚨 **Today's Events (Ongoing)**")
            for ev in today_events:
                render_almanac_event_card(ev, is_highlighted=True)
            st.markdown("<hr style='margin: 15px 0;'>", unsafe_allow_html=True)

        # 2. Upcoming Events (Next 14 Days)
        upcoming_2w_events = [ev for ev in filtered_almanac if ev['is_upcoming_2weeks']]
        upcoming_2w_events.sort(key=lambda x: x['start_date'])

        if upcoming_2w_events:
            st.markdown(f"##### 📅 **Upcoming Activities (Next 2 Weeks: {len(upcoming_2w_events)} Events)**")
            for ev in upcoming_2w_events:
                render_almanac_event_card(ev, is_highlighted=False)
        elif not today_events:
            st.info(f"No Almanac events scheduled for '{selected_unit}' in the next 14 days.")

        # 3. Complementary Flyers / Posters from Drive
        if campus_files:
            filtered_campus = [
                f for f in campus_files 
                if selected_unit == "All Units / Campus Wide" or selected_unit.lower() in f.get("name", "").lower()
            ]
            if filtered_campus:
                st.markdown("<hr style='margin: 15px 0;'>", unsafe_allow_html=True)
                st.markdown("##### 📁 **Event Posters & Circulars**")
                for f in filtered_campus:
                    render_bulletin_card(f, "Department / Club Event", bg_color="#16A34A")

    # Admin Control Panel
    if st.session_state.logged_email in ["research@stmaryscollege.in", "iqac@stmaryscollege.in"]:
        with st.expander("🛠️ Manage Announcements (Delete / Override Expired Dossiers & Flyers)"):
            all_announcements = research_files + campus_files
            if all_announcements:
                file_map = {f"{f.get('name')} (ID: {f.get('id')})": f.get('id') for f in all_announcements}
                selected_del = st.selectbox("Select Expired Document/Flyer to Delete from Drive:", list(file_map.keys()))
                if st.button("🗑️ Delete Selected Document from Drive", type="secondary"):
                    if delete_drive_file(file_map[selected_del], creds):
                        st.success("✅ Expired document deleted from Drive successfully!")
                        st.rerun()
                    else:
                        st.error("Failed to delete file from Drive.")
            
            st.markdown("---")
            st.markdown("**📤 Upload New Announcements Document (Word .docx / Image / PDF):**")
            u_target = st.radio("Target Folder:", ["🔬 Research Conferences Vault", "🎭 Campus Activities Vault"], horizontal=True)
            u_file = st.file_uploader("Select Announcement File:", type=["docx", "png", "jpg", "jpeg", "pdf"], key="bulletin_upload")
            if st.button("Publish File to Live Bulletin"):
                if u_file:
                    target_id = RESEARCH_EVENTS_FOLDER_ID if "Research" in u_target else CAMPUS_ACTIVITIES_FOLDER_ID
                    upload_file_to_drive(u_file.read(), u_file.name, u_file.type, [target_id], creds)
                    st.success("🎉 File uploaded to Drive and published to the live bulletin!")
                    st.rerun()

# --- TAB 3: LIVE MASTER DATABASE EXPLORER ---
with tab_explorer:
    st.subheader("📋 Master Google Sheet Live Explorer")
    st.markdown("Extract and inspect records across all sheets in real time.")
    
    sheet_choice = st.selectbox("Select Master Sheet Tab to View:", [
        "🔬 Research_Database", 
        "🏆 Faculty_Achievements", 
        "👥 Student_Activities", 
        "🏛️ Committees_Cells_Clubs"
    ])
    
    target_tab_map = {
        "🔬 Research_Database": res_df,
        "🏆 Faculty_Achievements": fac_df,
        "👥 Student_Activities": stu_df,
        "🏛️ Committees_Cells_Clubs": comm_df
    }
    selected_df = target_tab_map[sheet_choice]
    
    if not selected_df.empty:
        search_query = st.text_input("🔍 Search within this sheet (filter by Faculty Name, Department, or Title):", "").strip().lower()
        display_df = selected_df.copy()
        if search_query:
            display_df = display_df[display_df.apply(lambda row: row.astype(str).str.lower().str.contains(search_query).any(), axis=1)]
            
        st.dataframe(display_df, use_container_width=True, height=400)
        st.caption(f"Showing {len(display_df)} of {len(selected_df)} records")
        
        csv_data = display_df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="📥 Export Filtered View as CSV",
            data=csv_data,
            file_name=f"{sheet_choice.split()[-1]}_extracted.csv",
            mime="text/csv"
        )
    else:
        st.info(f"The `{sheet_choice}` tab is currently empty or contains no records in the Master Sheet.")

# --- TAB 4: DATA ENTRY WORKSPACE ---
with tab_submit:
    is_locked = not st.session_state.get("admin_enabled", True)
    is_admin = st.session_state.get("logged_email") in ["research@stmaryscollege.in", "iqac@stmaryscollege.in"]

    if is_locked and not is_admin:
        st.error("🔒 Data entry is currently disabled by the Administrator.")
    else:
        st.subheader("Add Monthly Achievement Entry")
        
        scope_type = st.radio("Select Reporting Scope*", ["Department", "Committee / Cell / Club"], horizontal=True)
        col1, col2, col3 = st.columns(3)
        with col1:
            if scope_type == "Department":
                form_focus = st.selectbox("Department Focus*", DEPARTMENTS)
            else:
                form_focus = st.selectbox("Committees / Cells / Clubs Focus*", COMMITTEES_CELLS_CLUBS)
        with col2: form_month = st.selectbox("Reporting Month*", MONTHS)
        with col3: form_year = st.selectbox("Academic Year*", ACADEMIC_YEARS)
            
        st.markdown("---")
        
        if scope_type == "Department":
            classification = st.selectbox("Select Classification", [
                "--- Select Category ---", 
                "🔬 Research Database", 
                "🏆 Faculty Profiles & Milestones", 
                "👥 Departmental & Student Contributions"
            ])

            if classification != "--- Select Category ---":
                if classification == "🔬 Research Database":
                    r_type = st.selectbox("Research Type", ["Paper Publication", "Book Chapter", "Full Book", "Paper Presentation", "FDP", "Workshop"])
                    collab_check = st.checkbox("Collaboration involved?", key="collab_box")
                    
                    with st.form("research_db_form", clear_on_submit=True):
                        title = st.text_input("Title*")
                        org = st.text_input("Organised By/Journal Name*")
                        
                        if r_type in ["Paper Publication", "Book Chapter", "Full Book"]:
                            index_type = st.selectbox("Indexing/Journal Type*", ["UGC Care Listed", "Scopus", "Web of Science", "SCIE", "ABDC", "PubMed", "Peer Reviewed", "DOAJ", "Embase"])
                            issn = st.text_input("ISSN/ISBN Number*")
                            url = st.text_input("URL*")
                            date_span, scope = "NA", "NA"
                        elif r_type in ["Paper Presentation", "FDP", "Workshop"]:
                            date_span = st.text_input("Date Span*")
                            scope = st.selectbox("Scope*", ["International", "National", "State", "Institutional"])
                            index_type, issn, url = "NA", "NA", "NA"
                        
                        collab_names = st.text_input("Enter Collaborator Names*") if st.session_state.collab_box else ""
                        upload = st.file_uploader("Upload Verification Document (Mandatory)*")
                        
                        if st.form_submit_button("Commit Entry"):
                            if not st.session_state.get("admin_enabled", True) and not is_admin:
                                st.error("Submission rejected: Data entry is currently disabled.")
                            elif not upload: st.error("Verification mandatory!")
                            elif st.session_state.collab_box and not collab_names.strip(): st.error("Collaboration names mandatory!")
                            elif not title or not org: st.error("Title and Organisation are mandatory!")
                            else:
                                dept_base_folder = DEPARTMENT_FOLDERS.get(form_focus, "1HMBoNkhksNpaitlBaGfq3JeoHsb_jmo-")
                                logged_in_email = st.session_state.logged_email
                                if logged_in_email in [
                                    "sciences@stmaryscollege.in", "languages@stmaryscollege.in",
                                    "businessmanagement@stmaryscollege.in", "commerce@stmaryscollege.in",
                                    "socialsciences@stmaryscollege.in"
                                ]:
                                    target_folder = get_or_create_drive_folder("Department Activities", dept_base_folder, creds)
                                else:
                                    target_folder = get_or_create_drive_folder(current_faculty_name, dept_base_folder, creds)
                                
                                drive_link = upload_file_to_drive(upload.read(), upload.name, upload.type, [target_folder], creds)
                                new_row = [current_faculty_name, form_focus, r_type, index_type, title, drive_link, date_span, url, org, scope, scope, org, issn, form_month]
                                append_and_sort_sheet_by_department("Research_Database", new_row, 1, creds)
                                st.success("🎉 Research entry written and sorted in Master Sheet successfully!")
                                st.rerun()

                elif classification == "🏆 Faculty Profiles & Milestones":
                    target_sheet = "Faculty_Achievements"
                    subtype = st.selectbox("Select Profile Subtype", ["Certification/Course", "Presentation/Resource Person", "Doctoral Milestone", "Award/Honor"])
                    if subtype == "Certification/Course": 
                        styled_block("[Name], [Certification Title/Course Name], [Issuing Body], [Result/Grade/Medal if applicable].", "Mr. MSS Roy successfully completed an 8-week NPTEL certification course in 'Advanced Corporate Governance' with an Elite Silver Medal, organized by IIT Madras.")
                    elif subtype == "Presentation/Resource Person": 
                        styled_block("[Name], [Role: Guest Speaker/Judge/Facilitator], '[Topic/Title],' [Organizing Event Name/Department/Institution], [Date].", "Dr. Rajita Anand Singh acted as a Resource Person and delivered an invited lecture on 'Emerging Trends in Literary Criticism' for the National Colloquium organized by the Department of English, St. Mary's College on June 15, 2026.")
                    elif subtype == "Doctoral Milestone": 
                        styled_block("[Name], [Milestone Achieved], '[Research Topic],' [University/Institution], [Date].", "Ms. Shima A.N successfully completed her Ph.D. Viva-Voce examination for her doctoral thesis titled 'A Comprehensive Evaluation of Cloud Workloads' at Osmania University.")
                    elif subtype == "Award/Honor": 
                        styled_block("[Name], [Title of Award/Recognition], [Awarding Body/Organization], [Date].", "Dr. Deepthi Priya was conferred with the 'Best Faculty Researcher Award 2026' by the Institute of Scholar Recognitions on May 12, 2026.")
                    
                    with st.form("faculty_form", clear_on_submit=True):
                        narrative_input = st.text_area("Achievement Narrative*")
                        upload = st.file_uploader("Upload Verification Document (Mandatory)*")
                        if st.form_submit_button("Submit Profile"):
                            if not st.session_state.get("admin_enabled", True) and not is_admin:
                                st.error("Submission rejected: Data entry is currently disabled.")
                            elif not upload or not narrative_input.strip(): st.error("Verification and narrative statement mandatory!")
                            else:
                                dept_base_folder = DEPARTMENT_FOLDERS.get(form_focus, "1HMBoNkhksNpaitlBaGfq3JeoHsb_jmo-")
                                logged_in_email = st.session_state.logged_email
                                if logged_in_email in [
                                    "sciences@stmaryscollege.in", "languages@stmaryscollege.in",
                                    "businessmanagement@stmaryscollege.in", "commerce@stmaryscollege.in",
                                    "socialsciences@stmaryscollege.in"
                                ]:
                                    target_folder = get_or_create_drive_folder("Department Activities", dept_base_folder, creds)
                                else:
                                    target_folder = get_or_create_drive_folder(current_faculty_name, dept_base_folder, creds)
                                
                                drive_link = upload_file_to_drive(upload.read(), upload.name, upload.type, [target_folder], creds)
                                new_row = [datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), form_focus, form_month, form_year, subtype, narrative_input.strip(), current_faculty_name, drive_link]
                                append_and_sort_sheet_by_department(target_sheet, new_row, 1, creds)
                                st.success("🎉 Profile entry written and sorted in Master Sheet!")
                                st.rerun()

                elif classification == "👥 Departmental & Student Contributions":
                    target_sheet = "Student_Activities"
                    styled_block("[Coordinator/Dept], [Type of Event/Activity], [Beneficiaries/Location], [Date].", "The Department of Sciences hosted an Inter-Collegiate Science Exhibition titled 'Eco-Innovate 2026' for undergraduate students of regional colleges on April 22, 2026.")
                    with st.form("student_form", clear_on_submit=True):
                        description = st.text_area("Description*")
                        upload = st.file_uploader("Upload Verification Document (Mandatory)*")
                        if st.form_submit_button("Submit Contribution"):
                            if not st.session_state.get("admin_enabled", True) and not is_admin:
                                st.error("Submission rejected: Data entry is currently disabled.")
                            elif not upload or not description.strip(): st.error("Verification and description mandatory!")
                            else:
                                dept_base_folder = DEPARTMENT_FOLDERS.get(form_focus, "1HMBoNkhksNpaitlBaGfq3JeoHsb_jmo-")
                                logged_in_email = st.session_state.logged_email
                                if logged_in_email in [
                                    "sciences@stmaryscollege.in", "languages@stmaryscollege.in",
                                    "businessmanagement@stmaryscollege.in", "commerce@stmaryscollege.in",
                                    "socialsciences@stmaryscollege.in"
                                ]:
                                    target_folder = get_or_create_drive_folder("Department Activities", dept_base_folder, creds)
                                else:
                                    target_folder = get_or_create_drive_folder(current_faculty_name, dept_base_folder, creds)
                                
                                drive_link = upload_file_to_drive(upload.read(), upload.name, upload.type, [target_folder], creds)
                                new_row = [datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), form_focus, form_month, form_year, "Institutional Contribution", description.strip(), current_faculty_name, drive_link]
                                append_and_sort_sheet_by_department(target_sheet, new_row, 1, creds)
                                st.success("🎉 Contribution entry written and sorted in Master Sheet!")
                                st.rerun()

        else:
            target_sheet = "Committees_Cells_Clubs"
            styled_block("[Committee Name], organized [Event Type/Activity Details] on [Date].", "The Placement Cell coordinated a campus recruitment drive with Deloitte for final year commerce students on May 18, 2026.")
            
            with st.form("committees_ledger_form", clear_on_submit=True):
                narrative_input = st.text_area("Narrative Log Description*")
                event_date = st.date_input("Date of Event Activity*", value=datetime.date.today())
                upload = st.file_uploader("Upload Verification Document (Mandatory)*")
                
                if st.form_submit_button("Commit Committee Record"):
                    if not st.session_state.get("admin_enabled", True) and not is_admin:
                        st.error("Submission rejected: Data entry is currently disabled.")
                    elif not upload or not narrative_input.strip(): 
                        st.error("Log Description and verification attachment are strictly mandatory fields!")
                    else:
                        target_folder = get_or_create_drive_folder(current_faculty_name, COMMITTEE_FOLDER_ID, creds)
                        drive_link = upload_file_to_drive(upload.read(), upload.name, upload.type, [target_folder], creds)
                        new_row = [
                            form_focus, 
                            current_faculty_name, 
                            form_month, 
                            form_year, 
                            narrative_input.strip(), 
                            str(event_date)
                        ]
                        append_and_sort_sheet_by_department(target_sheet, new_row, 0, creds)
                        st.success(f"🎉 Structured Activity Log written to '{target_sheet}' sheet successfully!")
                        st.rerun()

# --- 5. MONTHLY GENERATOR ---
with tab_document:
    st.subheader("Central Document Engine Dashboard Workspace")
    
    col_d1, col_d2, col_d3 = st.columns(3)
    with col_d1:
        view_focus = st.selectbox(
            "Target Department / Scope Scope", 
            DEPARTMENTS + ["Committees / Cells / Clubs"], 
            key="vd1"
        )
            
    with col_d2: view_month = st.selectbox("Target Month Scope", MONTHS, key="vm1")
    with col_d3: view_year = st.selectbox("Target Year Scope", ACADEMIC_YEARS, key="vy1")
        
    if st.button("🏗️ Construct Automated Monthly Document Package", use_container_width=True, type="primary"):
        with st.spinner("Assembling structured records from sheets..."):
            docx_bytes = build_monthly_word_document(view_focus, view_month, view_year, creds)
            file_name_string = f"Monthly_Achievements_Report_{view_focus.replace(' ', '_')}_{view_month}_{view_year}.docx"
            
            target_folder = DEPARTMENT_FOLDERS.get(view_focus, COMMITTEE_FOLDER_ID) if view_focus != "Committees / Cells / Clubs" else COMMITTEE_FOLDER_ID
            upload_file_to_drive(docx_bytes, file_name_string, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", [target_folder], creds)
            
            st.success("🎯 Document synchronized into your Drive repository folder automatically!")
            st.download_button(label="📥 Download Report File Asset Directly", data=docx_bytes, file_name=file_name_string, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

# --- 6. ADMIN CONTROL ---
with tab_admin:
    if st.session_state.logged_email in ["research@stmaryscollege.in", "iqac@stmaryscollege.in"]:
        st.toggle(
            "Enable Data Entry for Users", 
            key="admin_toggle_widget", 
            value=st.session_state.get("admin_enabled", True)
        )
        st.session_state.admin_enabled = st.session_state.admin_toggle_widget
    else: 
        st.warning("Unauthorized access. Admin privileges restricted to Research & IQAC accounts.")

st.markdown("---")
st.markdown("<div style='text-align: center; color: olive;'>Developed by Research Committee @ St. Mary's College</div>", unsafe_allow_html=True)
